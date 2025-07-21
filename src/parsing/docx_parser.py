# DOCX-specific parser to handle tables and formatted text
from docx import Document
from pathlib import Path
from typing import Optional
import re

'''def read_docx_file(file_path):
    """
    Reads text from a .docx file including paragraphs and table cells.
    
    Args:
        file_path: Path to the .docx file.
        
    Returns:
        All extracted text as a single string, or None if the file is empty or doesn't exist.
    """
    try:
        if not Path(file_path).exists():
            print("File not found.")
            return None

        doc = Document(file_path)
        all_text = []'''

def extract_text_from_docx(docx_path: str) -> Optional[str]:
    """
    Extract text from DOCX including tables
    
    Args:
        docx_path: Path to DOCX file
        
    Returns:
        Combined text from paragraphs and tables
    """
    try:
        if not Path(docx_path).exists():
            raise FileNotFoundError(f"DOCX file not found: {docx_path}")
            
        doc = Document(docx_path)
        text_parts = []
        
        # 1. Extract paragraphs
        for p in doc.paragraphs:
            if p.text.strip():
                text_parts.append(p.text)
        # text_parts.extend(p.text for p in doc.paragraphs if p.text.strip())
        
        # 2. Extract table content
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text_parts.append(cell.text)
        
        # 3. Combine and clean
        combined = '\n'.join(text_parts)
        return combined if combined.strip() else None
        
    except Exception as e:
        print(f"[python-docx] Extraction error: {e}")
        return None