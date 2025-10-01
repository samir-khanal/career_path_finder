from pathlib import Path
from fpdf import FPDF
import docx

def create_simple_pdf():
    """Create a simple PDF for testing"""
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)

    pdf.cell(200, 10, txt="John Doe - Data Scientist", ln=1, align='C')
    pdf.ln(10)

    pdf.cell(200, 10, txt="SKILLS:", ln=1)
    pdf.multi_cell(0, 10, txt="Python, Machine Learning, SQL, Data Visualization, Statistics, Deep Learning")
    pdf.ln(5)

    pdf.cell(200, 10, txt="EDUCATION:", ln=1)
    pdf.multi_cell(0, 10, txt="MS in Computer Science - Tech University (2022-2024)")
    pdf.ln(5)

    pdf.cell(200, 10, txt="EXPERIENCE:", ln=1)
    pdf.multi_cell(0, 10, txt="Data Scientist at AI Company (2024-Present)\n- Developed ML models\n- Analyzed big data")

    output_path = Path("tests/test_data/pdfs/simple.pdf")
    output_path.parent.mkdir(parents=True, exist_ok=True)
    pdf.output(str(output_path))
    print(f"✅ Created: {output_path}")

def create_table_pdf():
    """Create a PDF with tables for testing"""
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)

# Personal Info
    pdf.cell(200, 10, txt="JOHN SMITH - SENIOR DATA SCIENTIST", ln=1, align='C')
    pdf.cell(200, 10, txt="john.smith@email.com | (555) 987-6543 | linkedin.com/in/johnsmith", ln=1, align='C')
    pdf.ln(15)

    # SKILLS TABLE
    pdf.set_font("Arial", 'B', 14)
    pdf.cell(200, 10, txt="TECHNICAL SKILLS", ln=1)
    pdf.set_font("Arial", size=11)
    #Skills Table 1: Programming Languages
    pdf.cell(200, 8, txt="Programming Languages:", ln=1)
    skills_data = [
        ["Python", "5 years", "Expert"],
        ["SQL", "4 years", "Advanced"], 
        ["Machine Learning", "3 years", "Expert"],
        ["Tableau", "2 years", "Intermediate"]
    ]
    
    col_width = 60
    row_height = 8
    for row in skills_data:
        for item in row:
            pdf.cell(col_width, row_height, txt=str(item), border=1)
        pdf.ln(row_height)
    pdf.ln(5)

    # Skills Table 2: Data Science Tools
    pdf.cell(200, 8, txt="Data Science Tools:", ln=1)
    data2 = [
        ["Tool", "Category", "Proficiency"],
        ["TensorFlow", "ML Framework", "Expert"],
        ["Tableau", "Visualization", "Advanced"],
        ["Scikit-learn", "ML Library", "Expert"],
        ["PyTorch", "Deep Learning", "Intermediate"]
    ]

    for row in data2:
        for item in row:
            pdf.cell(col_width, row_height, txt=str(item), border=1)
        pdf.ln(row_height)

    pdf.ln(15)

    # EDUCATION TABLE
    pdf.set_font("Arial", 'B', 14)
    pdf.cell(200, 10, txt="EDUCATION", ln=1)
    pdf.set_font("Arial", size=11)

    edu_data = [
        ["Degree", "Institution", "Year", "GPA"],
        ["PhD Data Science", "Tech University", "2022-2024", "3.9/4.0"],
        ["MS Computer Science", "State University", "2018-2020", "3.8/4.0"],
        ["BS Mathematics", "City College", "2014-2018", "3.7/4.0"]
    ]

    col_widths = [50, 50, 30, 20]

    for row in edu_data:
        for i, item in enumerate(row):
            pdf.cell(col_widths[i], 8, txt=str(item), border=1)
        pdf.ln(8)

    pdf.ln(15)

    # EXPERIENCE TABLE
    pdf.set_font("Arial", 'B', 14)
    pdf.cell(200, 10, txt="PROFESSIONAL EXPERIENCE", ln=1)
    pdf.set_font("Arial", size=11)

    exp_data = [
        ["Position", "Company", "Duration", "Technologies"],
        ["Senior Data Scientist", "AI Innovations", "2022-Present", "Python, TensorFlow, AWS"],
        ["Data Scientist", "Tech Solutions", "2020-2022", "SQL, Tableau, Scikit-learn"],
        ["Data Analyst", "DataWorks Inc", "2018-2020", "Python, Excel, SQL"]
    ]

    col_widths_exp = [50, 40, 30, 60]

    for row in exp_data:
        for i, item in enumerate(row):
            pdf.cell(col_widths_exp[i], 8, txt=str(item), border=1)
        pdf.ln(8)

    # Save to file
    output_path = Path("tests/test_data/pdfs/tables.pdf")
    output_path.parent.mkdir(parents=True, exist_ok=True)
    pdf.output(str(output_path))
    print(f"✅ Created: {output_path}")

def create_simple_docx():
    """Create a simple DOCX for testing"""
    doc = docx.Document()
    doc.add_heading('John Doe - Data Scientist', 0)
    
    doc.add_heading('SKILLS', level=1)
    doc.add_paragraph('Python, Machine Learning, SQL, Data Visualization, Statistics, Deep Learning')
    
    doc.add_heading('EDUCATION', level=1)
    doc.add_paragraph('MS in Computer Science - Tech University (2022-2024)')
    
    doc.add_heading('EXPERIENCE', level=1)
    doc.add_paragraph('Data Scientist at AI Company (2024-Present)')
    doc.add_paragraph('- Developed ML models')
    doc.add_paragraph('- Analyzed big data')

    output_path = Path("tests/test_data/docs/simple.docx")
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    print(f"✅ Created: {output_path}")

def create_all_test_files():
    """Create all test files"""
    print("🛠️ Creating test files...")
    create_simple_pdf()
    create_table_pdf()
    create_simple_docx()
    print("✅ All test files created!")

if __name__ == "__main__":
    create_all_test_files()