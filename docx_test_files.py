# create_docx_test_files.py
from docx import Document
from pathlib import Path

def create_proper_docx_files():
    """Create actual .docx files for testing"""
    docs_dir = Path("tests/test_data/docs")
    docs_dir.mkdir(parents=True, exist_ok=True)
    
    # 1. simple.docx - Basic resume
    doc = Document()
    doc.add_heading('JOHN DOE - DATA SCIENTIST', 0)
    
    doc.add_heading('SKILLS', level=1)
    doc.add_paragraph('Python, SQL, Machine Learning, Data Analysis, Tableau')
    
    doc.add_heading('EDUCATION', level=1)
    doc.add_paragraph('Bachelor of Computer Science - University of Technology (2020-2024)')
    
    doc.add_heading('EXPERIENCE', level=1)
    doc.add_paragraph('Data Analyst - ABC Corp (2022-Present)')
    doc.add_paragraph('- Analyzed data using Python')
    doc.add_paragraph('- Created reports and dashboards')
    
    doc.save(docs_dir / "simple.docx")
    print("✅ Created simple.docx")
    
    # 2. tables.docx - Resume with tables
    doc2 = Document()
    doc2.add_heading('JANE SMITH - DATA ANALYST', 0)
    
    doc2.add_heading('TECHNICAL SKILLS', level=1)
    
    # Create skills table
    table = doc2.add_table(rows=4, cols=3)
    table.style = 'Table Grid'
    
    # Header row
    table.cell(0, 0).text = 'Category'
    table.cell(0, 1).text = 'Skills'
    table.cell(0, 2).text = 'Proficiency'
    
    # Data rows
    table.cell(1, 0).text = 'Programming'
    table.cell(1, 1).text = 'Python, SQL, JavaScript'
    table.cell(1, 2).text = 'Advanced'
    
    table.cell(2, 0).text = 'Data Science'
    table.cell(2, 1).text = 'Machine Learning, Tableau'
    table.cell(2, 2).text = 'Expert'
    
    table.cell(3, 0).text = 'Tools'
    table.cell(3, 1).text = 'Git, Docker, AWS'
    table.cell(3, 2).text = 'Intermediate'
    
    doc2.add_heading('EDUCATION', level=1)
    doc2.add_paragraph('Master of Data Science - Data University (2022-2024)')
    
    doc2.add_heading('EXPERIENCE', level=1)
    doc2.add_paragraph('Data Scientist - Tech Innovations Inc. (2024-Present)')
    
    doc2.save(docs_dir / "tables.docx")
    print("✅ Created tables.docx")

if __name__ == "__main__":
    create_proper_docx_files()
    print("🎉 All DOCX test files created successfully!")